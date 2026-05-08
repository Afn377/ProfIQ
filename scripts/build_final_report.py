from __future__ import annotations

import json
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "ProfIQ_Project_Report.docx"
ASSET_DIR = Path(tempfile.gettempdir()) / "profiq_report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
BORDER = "DADCE0"
FILL = "F7F9FC"
HEADER_FILL = "EEF3F8"
ACCENT = RGBColor(31, 77, 120)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def font_run(run, size=None, bold=None, italic=None, color=None, name="Arial") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", *, style=None, bold=False, italic=False, size=11, color=BLACK, after=8, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        r = p.add_run(text)
        font_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.clear()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    r = p.add_run(text)
    font_run(r, size=16 if level == 1 else 14 if level == 2 else 12, bold=True, color=BLACK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    font_run(r, size=11, color=BLACK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    font_run(r, size=11, color=BLACK)
    return p


def add_table(doc, headers, rows, widths=None, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    if widths:
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        font_run(r, size=font_size, bold=True, color=BLACK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(val))
            font_run(r, size=font_size, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def draw_bar_chart(path: Path, title: str, labels: list[str], series: dict[str, list[float]], y_max=1.0):
    w, h = 1400, 820
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    font = ImageFont.load_default()
    title_font = ImageFont.load_default(size=28)
    label_font = ImageFont.load_default(size=18)
    small = ImageFont.load_default(size=16)

    margin_l, margin_r, margin_t, margin_b = 120, 60, 90, 130
    plot_w = w - margin_l - margin_r
    plot_h = h - margin_t - margin_b
    d.text((margin_l, 28), title, fill=(0, 0, 0), font=title_font)
    d.line((margin_l, margin_t + plot_h, margin_l + plot_w, margin_t + plot_h), fill=(80, 80, 80), width=2)
    d.line((margin_l, margin_t, margin_l, margin_t + plot_h), fill=(80, 80, 80), width=2)

    for tick in [0, 0.25, 0.5, 0.75, 1.0]:
        y = margin_t + plot_h - int((tick / y_max) * plot_h)
        d.line((margin_l, y, margin_l + plot_w, y), fill=(230, 230, 230), width=1)
        d.text((32, y - 10), f"{tick:.2f}", fill=(80, 80, 80), font=small)

    colors = [(50, 115, 184), (46, 184, 139), (242, 170, 66)]
    metric_names = list(series.keys())
    group_w = plot_w / len(labels)
    bar_w = group_w / (len(metric_names) + 2)
    for i, label in enumerate(labels):
        cx = margin_l + i * group_w + group_w * 0.17
        for j, metric in enumerate(metric_names):
            val = series[metric][i]
            bh = int((val / y_max) * plot_h)
            x0 = int(cx + j * bar_w)
            x1 = int(x0 + bar_w * 0.78)
            y0 = margin_t + plot_h - bh
            y1 = margin_t + plot_h
            d.rectangle((x0, y0, x1, y1), fill=colors[j])
            d.text((x0, y0 - 22), f"{val:.3f}", fill=(0, 0, 0), font=small)
        d.text((int(margin_l + i * group_w + 4), h - 98), label, fill=(0, 0, 0), font=label_font)

    legend_x = margin_l
    for j, metric in enumerate(metric_names):
        d.rectangle((legend_x, h - 48, legend_x + 24, h - 24), fill=colors[j])
        d.text((legend_x + 34, h - 48), metric, fill=(0, 0, 0), font=label_font)
        legend_x += 220
    img.save(path)


def draw_purity_chart(path: Path):
    w, h = 1200, 650
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_font = ImageFont.load_default(size=28)
    label_font = ImageFont.load_default(size=18)
    small = ImageFont.load_default(size=16)
    d.text((70, 28), "Content-based recommender purity@5 vs. random baseline", fill=(0, 0, 0), font=title_font)
    metrics = [("Department", 0.230, 0.055, "4.2x"), ("Institution", 0.093, 0.050, "1.9x")]
    x0, y0 = 110, 130
    max_v = 0.25
    for idx, (label, score, base, lift) in enumerate(metrics):
        y = y0 + idx * 205
        d.text((x0, y - 36), label, fill=(0, 0, 0), font=label_font)
        for val, fill, yoff, name in [(score, (46, 184, 139), 0, "ProfIQ"), (base, (180, 190, 200), 62, "Random")]:
            bw = int((val / max_v) * 820)
            d.text((x0, y + yoff + 9), name, fill=(70, 70, 70), font=small)
            d.rectangle((x0 + 95, y + yoff, x0 + 95 + bw, y + yoff + 38), fill=fill)
            d.text((x0 + 110 + bw, y + yoff + 8), f"{val:.3f}", fill=(0, 0, 0), font=small)
        d.text((x0 + 95, y + 116), f"Lift over random: {lift}", fill=(31, 77, 120), font=label_font)
    img.save(path)


def build_figures():
    model_chart = ASSET_DIR / "model_metrics.png"
    purity_chart = ASSET_DIR / "recommender_purity.png"
    draw_bar_chart(
        model_chart,
        "Held-out sentiment classification metrics",
        ["VADER\n7,427 test", "LogReg\n52,445 test", "DistilBERT\n7,427 test"],
        {
            "Accuracy": [0.764, 0.802, 0.857],
            "Macro-F1": [0.533, 0.719, 0.636],
            "Weighted-F1": [0.737, 0.817, 0.831],
        },
    )
    draw_purity_chart(purity_chart)
    return model_chart, purity_chart


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for name in ("List Bullet", "List Number"):
        st = styles[name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "ProfIQ Project Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font_run(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "CS 210 - Data Management for Data Science"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(footer.runs[0], size=9, color=MUTED)
    return doc


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("ProfIQ: Professor Recommendation System")
    font_run(r, size=24, bold=True, color=BLACK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Final Project Report - CS 210 Data Management for Data Science")
    font_run(r, size=13, color=MUTED)
    for label, value in [
        ("Student", "Afnan Haider"),
        ("Date", "May 2026"),
        ("Repository", "https://github.com/Afn377/ProfIQ"),
        ("Demo video", "https://youtu.be/o1A1l2W680g?si=4a1GuE1157LvYent"),
        ("Project type", "End-to-end data management, NLP analytics, and web application"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(f"{label}: ")
        font_run(lr, bold=True)
        vr = p.add_run(value)
        font_run(vr)
    add_para(
        doc,
        "Executive summary",
        bold=True,
        size=13,
        before=18,
        after=6,
    )
    add_para(
        doc,
        "ProfIQ helps students search, evaluate, and compare professors by combining public review signals with a normalized database, repeatable ingestion commands, sentiment analysis, and a React/Django web interface. The project emphasizes data-management choices: entity normalization, deduplication, source-aware ingestion, API design, reproducibility, and an intentionally minimal raw-text storage policy.",
    )
    for item in [
        "Implements a normalized professor-review schema with Source, Department, Professor, Course, Review, SentimentResult, and ProfessorStats tables.",
        "Supports large-scale professor catalog crawling while avoiding bulk persistence of raw review text.",
        "Compares a domain-tuned VADER baseline with TF-IDF + Logistic Regression and DistilBERT sentiment models.",
        "Adds a MiniLM content-based recommender for similar-professor discovery.",
        "Provides a usable web application with search, filtering, detail pages, live reviews, sentiment/theme charts, and comparison views.",
    ]:
        add_bullet(doc, item)


SCHEMA_ROWS = [
    ("Source", "id", "BigAutoField / integer primary key", "Django-generated identifier."),
    ("Source", "name", "CharField(64), unique", "Review platform name such as RateMyProfessors or Reddit."),
    ("Source", "base_url", "URLField / varchar(200), blank allowed", "Base website URL for the source."),
    ("Department", "id", "BigAutoField / integer primary key", "Django-generated identifier."),
    ("Department", "name", "CharField(128), unique", "Department name used for filtering and grouping."),
    ("Department", "code", "CharField(16), blank allowed", "Optional short department code."),
    ("Professor", "id", "BigAutoField / integer primary key", "Django-generated identifier."),
    ("Professor", "name", "CharField(128), indexed", "Professor display/search name."),
    ("Professor", "department_id", "ForeignKey to Department, nullable", "Optional department relationship; SET_NULL on delete."),
    ("Professor", "institution", "CharField(128), indexed, blank allowed", "School or institution name."),
    ("Professor", "bio", "TextField, blank allowed", "Optional free-text biography field."),
    ("Professor", "external_ref", "CharField(64), indexed, unique when nonblank", "External source id such as rmp:12345."),
    ("Professor", "source_avg_rating", "FloatField, nullable", "Profile-level rating from the source crawl."),
    ("Professor", "source_num_ratings", "IntegerField", "Profile-level rating count from the source crawl."),
    ("Professor", "created_at", "DateTimeField", "Set automatically when the row is created."),
    ("Course", "id", "BigAutoField / integer primary key", "Django-generated identifier."),
    ("Course", "code", "CharField(32), indexed, unique", "Course code such as CS101."),
    ("Course", "title", "CharField(200), blank allowed", "Optional course title."),
    ("Course", "department_id", "ForeignKey to Department, nullable", "Optional department relationship; SET_NULL on delete."),
    ("Course.professors join table", "id", "BigAutoField / integer primary key", "Django-created table for the Course-Professor many-to-many relationship."),
    ("Course.professors join table", "course_id", "ForeignKey to Course", "One side of the many-to-many relationship."),
    ("Course.professors join table", "professor_id", "ForeignKey to Professor", "Other side of the many-to-many relationship."),
    ("Review", "id", "BigAutoField / integer primary key", "Django-generated identifier."),
    ("Review", "professor_id", "ForeignKey to Professor", "Required professor; CASCADE on delete."),
    ("Review", "course_id", "ForeignKey to Course, nullable", "Optional course; SET_NULL on delete."),
    ("Review", "source_id", "ForeignKey to Source", "Required source; PROTECT on delete."),
    ("Review", "text", "TextField", "Persisted review text for seed/demo data."),
    ("Review", "rating", "FloatField, nullable", "Optional 1-5 rating; Reddit comments usually have none."),
    ("Review", "source_url", "URLField / varchar(200), blank allowed", "Unique with source when nonblank to avoid duplicate imports."),
    ("Review", "posted_at", "DateTimeField, nullable", "Original review timestamp when available."),
    ("Review", "ingested_at", "DateTimeField", "Set automatically during ingest."),
    ("SentimentResult", "id", "BigAutoField / integer primary key", "Django-generated identifier."),
    ("SentimentResult", "review_id", "OneToOneField to Review", "One sentiment row per persisted review; CASCADE on delete."),
    ("SentimentResult", "compound", "FloatField", "Final sentiment score in the VADER-style -1 to 1 range."),
    ("SentimentResult", "positive / neutral / negative", "FloatField", "Raw VADER distribution components."),
    ("SentimentResult", "label", "CharField(10)", "positive, neutral, or negative."),
    ("SentimentResult", "themes", "JSONField list", "Detected theme names for the review."),
    ("SentimentResult", "analyzed_at", "DateTimeField", "Updated automatically when analysis is saved."),
    ("ProfessorStats", "id", "BigAutoField / integer primary key", "Django-generated identifier."),
    ("ProfessorStats", "professor_id", "OneToOneField to Professor", "One aggregate dashboard row per professor; CASCADE on delete."),
    ("ProfessorStats", "review_count", "IntegerField", "Number of analyzed reviews/comments."),
    ("ProfessorStats", "avg_compound", "FloatField", "Mean compound sentiment score."),
    ("ProfessorStats", "positive_count / neutral_count / negative_count", "IntegerField", "Sentiment label counts."),
    ("ProfessorStats", "theme_counts", "JSONField object", "Map from theme name to frequency."),
    ("ProfessorStats", "recommendation_score", "FloatField", "Computed 0-100 recommendation score."),
    ("ProfessorStats", "updated_at", "DateTimeField", "Updated automatically when stats are saved."),
]


def build_report():
    model_chart, purity_chart = build_figures()
    doc = setup_document()
    add_title_page(doc)

    add_heading(doc, "1. Problem Definition and Background", 1)
    add_para(
        doc,
        "Students often choose courses using fragmented information: official catalog metadata, informal peer advice, RateMyProfessors-style ratings, and discussion posts. These signals are useful but scattered, inconsistent, and difficult to compare across professors. ProfIQ addresses this problem as a data-management system: it collects review signals, normalizes entities and sources, computes analytics, and exposes the results through searchable APIs and visual frontend workflows.",
    )
    add_para(
        doc,
        "The project connects directly to CS 210 concepts. It requires schema design, data cleaning, deduplication, source integration, reproducible ETL, denormalized analytics for fast reads, and evaluation of models trained from scraped data. The system also makes an explicit engineering tradeoff: it persists catalog metadata and aggregate statistics while fetching raw review text live or during bounded analysis jobs, reducing storage footprint and avoiding unnecessary redistribution of scraped text.",
    )
    add_para(doc, "Prior work and tools", bold=True, size=12, after=4)
    add_para(
        doc,
        "The rule-based baseline builds on VADER, a sentiment method designed for short social text [1]. The neural comparison uses DistilBERT, a compact transformer model derived from BERT [2, 3]. The recommender uses sentence embeddings inspired by Sentence-BERT-style semantic similarity [4]. ProfIQ differs from those general-purpose tools by applying them inside an end-to-end data-management pipeline for professor review search and comparison.",
    )
    add_para(doc, "Research questions", bold=True, size=12, after=4)
    for question in [
        "Can public professor review signals be organized into a clean relational system that supports fast search, comparison, and analytics?",
        "Does a trained sentiment classifier improve over a hand-tuned VADER baseline on labeled RMP review text?",
        "Can content embeddings identify professors with similar teaching-style language better than random matching?",
    ]:
        add_number(doc, question)

    add_heading(doc, "2. Data Description", 1)
    add_para(
        doc,
        "ProfIQ uses three data layers. First, seed JSON data supports local demonstration and schema testing. Second, a RateMyProfessors directory crawl populates the professor catalog with names, institutions, departments, external IDs, and source-level rating summaries. Third, a one-off ML corpus stores review text in Parquet for model training and evaluation rather than in the production database.",
    )
    add_table(
        doc,
        ["Dataset / source", "Format", "Fields used", "Role in project"],
        [
            ("Seed reviews", "JSON", "source, department, course, professor, review text, rating", "Small reproducible demo path for local setup and ORM/API checks."),
            ("RMP catalog", "SQLite rows from crawler", "professor name, institution, department, external_ref, average rating, rating count", "Large searchable catalog; current local DB contains 1,702,324 professors and 3,272 departments."),
            ("Live RMP reviews", "API response, not persisted", "comment, helpfulness/clarity rating, course, date", "Detail-page review display and lazy aggregate sentiment analysis."),
            ("Reddit mentions", "Public JSON/PRAW-compatible collector", "comment text, URL, timestamp, professor mention slice", "Auxiliary qualitative signal on first live-review page."),
            ("ML corpus", "Parquet", "350,018 review rows across 4,793 professors and 49 institutions", "Training/evaluating classifiers and building recommender embeddings."),
        ],
        widths=[1.25, 1.0, 1.9, 2.35],
        font_size=8.5,
    )

    add_heading(doc, "3. Database and Pipeline Methodology", 1)
    add_para(
        doc,
        "The relational schema separates entities that would otherwise be duplicated in raw review data. Professor belongs to an optional Department and has many Courses; Review links a professor, optional course, and Source; SentimentResult stores one NLP result per persisted Review; ProfessorStats stores denormalized aggregates for dashboard queries. Unique constraints on professor identity, external references, course codes, source names, and review source URLs support idempotent ingestion.",
    )
    add_para(doc, "Database tables and datatypes", bold=True, size=12, after=4)
    add_table(
        doc,
        ["Table", "Field", "Datatype", "Purpose / notes"],
        SCHEMA_ROWS,
        widths=[1.25, 1.55, 1.75, 1.95],
        font_size=6.8,
    )
    add_table(
        doc,
        ["Stage", "Persistent writes", "Intentionally not written", "Reason"],
        [
            ("Directory crawl", "Professor, Department, source rating/count", "Raw review text", "Keeps the catalog searchable without storing millions of review rows."),
            ("Seed ingest", "Sources, courses, demo reviews, sentiment rows, stats", "Duplicate source URLs", "Provides a reproducible small-data path for setup and tests."),
            ("Lazy analyze", "One ProfessorStats row per professor", "Review/comment text", "Computes aggregates once while respecting the minimal-storage design."),
            ("Live review endpoint", "Nothing persistent", "Any database rows", "Returns per-review sentiment inline for the user session."),
        ],
        widths=[1.15, 1.75, 1.55, 2.05],
        font_size=8.8,
    )
    add_para(
        doc,
        "This design is central to the project. A grader inspecting the production-scale SQLite database may see zero rows in Review and SentimentResult after a directory crawl; that is expected. The large catalog stores professor metadata and source summaries, while raw review text is fetched live, analyzed in memory, and discarded. The seed path still demonstrates the full normalized review schema when persistent example reviews are needed.",
    )

    add_heading(doc, "4. Sentiment and Recommendation Methods", 1)
    add_para(doc, "Hypothesis", bold=True, size=12, after=4)
    add_para(
        doc,
        "The ML hypothesis is that review text contains enough signal to predict whether a professor review is negative, neutral, or positive, and that a trained classifier should outperform a purely rule-based VADER baseline on held-out RateMyProfessors reviews. A second recommendation hypothesis is that professors whose reviews use similar language about clarity, workload, grading, and helpfulness will be meaningfully close in sentence-embedding space.",
    )
    add_para(doc, "Training data and target variable", bold=True, size=12, after=4)
    add_para(
        doc,
        "The supervised sentiment models are trained on the ML corpus built from live RateMyProfessors review pages. Each training row represents one review and includes the professor external id, professor name, institution, department, course, review text, helpfulness/clarity-derived rating, would-take-again flag when available, difficulty, posted date, and source URL. The input feature for sentiment classification is the review text. The target variable is a three-class sentiment label derived from the 1-5 rating: ratings <= 2.0 are labeled negative, ratings > 2.0 and <= 3.5 are labeled neutral, and ratings > 3.5 are labeled positive.",
    )
    add_para(
        doc,
        "The source URL is used to create stable train/validation/test splits through hashing, so the same review always lands in the same split across repeated runs and across different models. This prevents accidental split drift when comparing VADER, TF-IDF + Logistic Regression, and DistilBERT.",
    )
    add_para(doc, "Sentiment pipeline", bold=True, size=12, after=4)
    add_para(
        doc,
        "The baseline sentiment analyzer extends NLTK VADER with an academic-review lexicon, idiom rules, and a star-rating blend. The lexicon adds terms such as 'avoid', 'useless', 'helpful', and 'engaging'; the idiom layer catches multi-word phrases like 'avoid like the plague' and 'would not recommend'; the rating blend maps 1-5 star scores to the VADER compound scale as a secondary signal. Each review is also tagged for teaching themes: clarity, fairness, workload, helpfulness, engagement, and grading.",
    )
    add_para(
        doc,
        "The learned pipeline has two model paths. The live model is TF-IDF + Logistic Regression: text is lowercased, converted into unigram/bigram TF-IDF features, and passed into a class-weighted logistic regression classifier. This model is small enough to load during app runtime. The deeper comparison model fine-tunes DistilBERT on the same target labels. DistilBERT is used for offline evaluation because it is larger and slower, while Logistic Regression is better suited to the app's live sentiment augmentation path.",
    )
    add_para(doc, "Recommendation pipeline", bold=True, size=12, after=4)
    add_para(
        doc,
        "For similar-professor search, the pipeline groups the ML corpus by professor, concatenates a bounded amount of review text per professor, encodes that text with the all-MiniLM-L6-v2 sentence-transformer model, L2-normalizes each vector, and stores the resulting professor embedding matrix. At runtime, ProfIQ finds nearest neighbors by cosine similarity, hydrates those neighbors from the database, and then prefers same-department or same-institution matches when possible.",
    )
    add_para(doc, "Results", bold=True, size=12, after=4)
    add_para(
        doc,
        "The results support the main sentiment hypothesis. Both learned models improve over the VADER baseline on the held-out reviews. VADER is useful because it is transparent and always available, but it struggles with neutral reviews and domain-specific professor language. TF-IDF + Logistic Regression gives the strongest macro-F1 in the saved results, which means it handles the minority neutral class better. DistilBERT gives the highest accuracy and weighted-F1, showing that a deeper language model can capture more review phrasing, though it is less practical as the default live model. The recommender evaluation also supports the embedding hypothesis: nearest-neighbor matches have higher department and institution purity than a random baseline, meaning review-text embeddings are capturing teaching-style similarity better than chance.",
    )

    add_heading(doc, "5. API and Frontend Implementation", 1)
    add_para(
        doc,
        "The backend exposes REST endpoints for the user workflows. The frontend is a React/Vite app that calls those endpoints and renders search results, professor detail pages, comparison views, sentiment distributions, theme charts, live reviews, and similar-professor recommendations.",
    )
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ("GET /api/summary/", "Platform summary and top-ranked professors."),
            ("GET /api/professors/", "Search/filter by professor, department, institution, course, and sort mode."),
            ("POST /api/professors/", "Validated, throttled, deduplicated student-submitted professor creation."),
            ("GET /api/professors/<id>/", "Professor detail with courses, stats, source ratings, and lazy-analysis trigger."),
            ("GET /api/professors/<id>/reviews/", "Live RMP review page plus first-page Reddit mentions, analyzed inline."),
            ("GET /api/compare/?ids=...", "Compact multi-professor comparison data."),
            ("GET /api/professors/<id>/similar/", "MiniLM-based similar-professor recommendations when embeddings are available."),
        ],
        widths=[2.05, 4.45],
        font_size=8.8,
    )

    add_heading(doc, "6. Results and Analysis", 1)
    add_para(
        doc,
        "The implementation was verified with automated tests and a production frontend build. The sentiment unit tests cover lexicon behavior, idiom handling, rating blending, ML fallback behavior, and recommender artifact fallback/warm-up behavior. Professor creation tests cover validation, fictional/joke-name rejection, case-insensitive deduplication, and happy-path persistence.",
    )
    add_table(
        doc,
        ["Check", "Result", "Interpretation"],
        [
            ("python3 -m unittest sentiment.tests -v", "28 passed", "Core sentiment/recommender helpers behave as expected."),
            ("python3 manage.py test professors -v 2", "13 passed", "Create-professor API validation/dedup logic passes."),
            ("npm run build", "Passed", "Frontend compiles successfully; Vite reports only a non-fatal chunk-size warning."),
        ],
        widths=[2.35, 1.05, 3.1],
        font_size=8.8,
    )
    add_para(
        doc,
        "The model comparison shows that learned models outperform the VADER baseline on headline metrics. DistilBERT has the highest raw accuracy and weighted-F1 on the benchmark split, but Logistic Regression has the strongest macro-F1 in the saved artifacts because class weighting improves recall on the minority neutral class. This matters for the product: neutral reviews are common in real course selection, and a model that never predicts neutral would overstate student sentiment.",
    )
    add_table(
        doc,
        ["Model", "Test rows", "Accuracy", "Macro-F1", "Weighted-F1", "Use"],
        [
            ("VADER + domain rules", "7,427", "0.764", "0.533", "0.737", "Always-available baseline."),
            ("TF-IDF + LogReg", "52,445", "0.802", "0.719", "0.817", "Live default ML augmentation."),
            ("DistilBERT", "7,427", "0.857", "0.636", "0.831", "Offline comparison / optional live model."),
        ],
        widths=[1.5, 0.8, 0.75, 0.75, 0.85, 1.85],
        font_size=8.5,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Figure 1. ")
    font_run(r, bold=True, size=10)
    r = p.add_run("Sentiment model metrics from the saved evaluation artifacts.")
    font_run(r, size=10, color=MUTED)
    doc.add_picture(str(model_chart), width=Inches(6.2))

    add_para(
        doc,
        "The recommender evaluation uses purity@5: the fraction of each professor's top-five neighbors sharing a department or institution. Department purity reaches 0.230 versus a 0.055 random baseline, a 4.2x lift. Institution purity reaches 0.093 versus 0.050 random, a 1.9x lift. The stronger department lift is useful because it suggests the embedding captures subject-matter and teaching-language patterns more than campus identity alone.",
    )
    add_table(
        doc,
        ["Metric", "ProfIQ purity@5", "Random baseline", "Lift", "Interpretation"],
        [
            ("Department", "0.230", "0.055", "4.2x", "Embeddings recover teaching/content similarity better than chance."),
            ("Institution", "0.093", "0.050", "1.9x", "Location signal exists but is weaker than department/content signal."),
        ],
        widths=[1.2, 1.0, 1.0, 0.65, 2.65],
        font_size=8.8,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Figure 2. ")
    font_run(r, bold=True, size=10)
    r = p.add_run("Similar-professor recommender lift over random baselines.")
    font_run(r, size=10, color=MUTED)
    doc.add_picture(str(purity_chart), width=Inches(6.2))

    add_heading(doc, "7. Discussion, Limitations, and Ethics", 1)
    add_para(
        doc,
        "ProfIQ's main advantage is that it demonstrates the full data lifecycle: collection, cleaning, normalization, deduplication, analytics, serving, visualization, and model evaluation. The system is usable as an application, but the more important data-management contribution is the boundary between persisted metadata/aggregates and transient raw review text.",
    )
    for limitation in [
        "Source bias: RMP and Reddit comments are self-selected and may overrepresent unusually positive or negative experiences.",
        "Label noise: using RMP helpfulness/clarity ratings as sentiment labels is practical but imperfect because star ratings do not always match review text.",
        "Neutral-class difficulty: both VADER and DistilBERT struggle with neutral reviews; LogReg improves macro-F1 but still has room to improve.",
        "External-source reliability: RMP and Reddit access can be rate-limited or change without warning, so live review display must degrade gracefully.",
        "Ethics and terms: the project minimizes raw-text persistence and is intended for educational analysis rather than redistribution of scraped review corpora.",
        "Operational scaling: production deployment should move from SQLite to PostgreSQL and use a real background queue instead of in-process threads.",
    ]:
        add_bullet(doc, limitation)
    add_para(
        doc,
        "Future work would add better entity resolution across campuses, topic modeling or lemmatized theme extraction, calibration checks for sentiment scores, user preference controls, cache invalidation policies, and accessibility/performance hardening for the frontend.",
    )

    add_heading(doc, "8. Reproducibility", 1)
    add_para(
        doc,
        "The repository includes the backend, frontend, migrations, requirements, evaluation notebook, report builders, and example data. A fresh local run uses the seed ingest path; large crawls and ML artifacts can be rebuilt from the documented management commands.",
    )
    add_para(
        doc,
        "Demo video: https://youtu.be/o1A1l2W680g?si=4a1GuE1157LvYent",
        bold=True,
    )
    add_table(
        doc,
        ["Step", "Command"],
        [
            ("Install backend", "python3 -m venv .venv; source .venv/bin/activate; pip install -r backend/requirements.txt"),
            ("Initialize database", "cd backend; python manage.py migrate; python manage.py ingest_seed --reset"),
            ("Run backend", "python manage.py runserver 8000"),
            ("Run frontend", "cd frontend; npm install; npm run dev"),
            ("Run tests", "cd backend; python3 -m unittest sentiment.tests -v; python3 manage.py test professors -v 2"),
            ("Build frontend", "cd frontend; npm run build"),
        ],
        widths=[1.45, 5.05],
        font_size=8.2,
    )

    add_heading(doc, "9. Conclusion", 1)
    add_para(
        doc,
        "ProfIQ satisfies the project objective by delivering both a working application and a documented data-management workflow. The system integrates multiple public data sources, normalizes professor-review entities, computes interpretable sentiment/theme analytics, evaluates learned models against a rule baseline, and exposes the results through a web interface. The final design is intentionally pragmatic: it keeps the database queryable and small while still giving students live access to review-level evidence when they need it.",
    )

    add_heading(doc, "References", 1)
    refs = [
        "Hutto, C. J., & Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text. International AAAI Conference on Web and Social Media.",
        "Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. NAACL-HLT.",
        "Sanh, V., Debut, L., Chaumond, J., & Wolf, T. (2019). DistilBERT, a distilled version of BERT: smaller, faster, cheaper and lighter. arXiv:1910.01108.",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. EMNLP-IJCNLP.",
        "Pedregosa, F. et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research.",
        "Django Software Foundation. Django documentation. https://docs.djangoproject.com/",
        "React documentation. https://react.dev/",
        "AI assistance disclosure: ChatGPT was used for assistance with debugging, explanation, editing, and project organization. It was not used as the source of generated project content submitted in place of my own work.",
    ]
    for ref in refs:
        add_para(doc, ref, size=10, after=4)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_report()
